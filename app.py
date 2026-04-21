from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
import tempfile
import os
import io

app = Flask(__name__)
CORS(app)

def add_page_numbers(section):
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

@app.route('/format', methods=['POST'])
def format_book():
    file = request.files['file']
    trim_width = float(request.form['trim_width'])
    trim_height = float(request.form['trim_height'])
    margin_top = float(request.form['margin_top'])
    margin_bottom = float(request.form['margin_bottom'])
    margin_inside = float(request.form['margin_inside'])
    margin_outside = float(request.form['margin_outside'])
    author_name = request.form['author_name']
    book_title = request.form['book_title']
    font_name = request.form.get('font_name', 'Garamond')
    font_size = float(request.form.get('font_size', 12))
    theme = request.form.get('theme', 'classic')

    themes = {
        'classic': {'font': 'Garamond', 'size': 11.5, 'heading_size': 14},
        'modern': {'font': 'Georgia', 'size': 12, 'heading_size': 16},
        'romance': {'font': 'Palatino Linotype', 'size': 11, 'heading_size': 13},
        'nonfiction': {'font': 'Times New Roman', 'size': 12, 'heading_size': 15},
        'minimalist': {'font': 'Book Antiqua', 'size': 11, 'heading_size': 12},
    }

    if theme in themes:
        font_name = themes[theme]['font']
        font_size = themes[theme]['size']
        heading_size = themes[theme]['heading_size']
    else:
        heading_size = 14

    doc = Document(file)
    section = doc.sections[0]

    section.page_width = Inches(trim_width)
    section.page_height = Inches(trim_height)
    section.top_margin = Inches(margin_top)
    section.bottom_margin = Inches(margin_bottom)
    section.left_margin = Inches(margin_inside)
    section.right_margin = Inches(margin_outside)

    # Replace scene breaks *** or --- with typographic ornament
    for para in doc.paragraphs:
        if para.text.strip() in ['***', '---', '* * *']:
            para.clear()
            run = para.add_run('✦')
            run.font.name = font_name
            run.font.size = Pt(font_size)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)
        if is_empty:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            continue

        if is_heading or is_title:
            para.paragraph_format.space_before = Pt(24)
            para.paragraph_format.space_after = Pt(6)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(heading_size)
            continue

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.first_line_indent = Inches(0.3)

        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run(f"{author_name}  |  {book_title}")
    run.font.name = font_name
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    add_page_numbers(section)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name='formatted_manuscript.docx'
    )

@app.route('/cover', methods=['POST'])
def generate_cover():
    front_image = request.files['front_image']
    back_image = request.files.get('back_image')
    spine_image = request.files.get('spine_image')
    trim_width = float(request.form['trim_width'])
    trim_height = float(request.form['trim_height'])
    page_count = int(request.form['page_count'])
    paper_type = request.form.get('paper_type', 'white')
    title = request.form.get('title', '')
    author = request.form.get('author', '')
    spine_color = request.form.get('spine_color', '#282828')

    if paper_type == 'cream':
        spine_width = page_count * 0.0025
    elif paper_type == 'color':
        spine_width = page_count * 0.002347
    else:
        spine_width = page_count * 0.002252

    bleed = 0.125
    total_width = (trim_width * 2) + spine_width + (bleed * 2)
    total_height = trim_height + (bleed * 2)

    dpi = 300
    total_width_px = int(total_width * dpi)
    total_height_px = int(total_height * dpi)
    trim_width_px = int(trim_width * dpi)
    spine_width_px = max(int(spine_width * dpi), 1)
    bleed_px = int(bleed * dpi)

    cover = Image.new('RGB', (total_width_px, total_height_px), (255, 255, 255))

    # BACK COVER
    if back_image:
        back_img = Image.open(back_image)
        back_img = back_img.resize((trim_width_px, total_height_px), Image.LANCZOS)
        cover.paste(back_img, (bleed_px, 0))
    else:
        back_fill = Image.new('RGB', (trim_width_px, total_height_px), (240, 240, 240))
        cover.paste(back_fill, (bleed_px, 0))

    # SPINE
    spine_x = bleed_px + trim_width_px
    if spine_image:
        sp_img = Image.open(spine_image)
        sp_img = sp_img.resize((spine_width_px, total_height_px), Image.LANCZOS)
        cover.paste(sp_img, (spine_x, 0))
    else:
        r = int(spine_color[1:3], 16)
        g = int(spine_color[3:5], 16)
        b = int(spine_color[5:7], 16)
        spine_fill = Image.new('RGB', (spine_width_px, total_height_px), (r, g, b))
        cover.paste(spine_fill, (spine_x, 0))

    # FRONT COVER
    front_img = Image.open(front_image)
    front_img = front_img.resize((trim_width_px, total_height_px), Image.LANCZOS)
    front_x = bleed_px + trim_width_px + spine_width_px
    cover.paste(front_img, (front_x, 0))

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg', dir=os.getcwd())
    cover.save(tmp.name, 'JPEG', quality=95, dpi=(300, 300))

    pdf_output = io.BytesIO()
    c = canvas.Canvas(pdf_output, pagesize=(total_width * inch, total_height * inch))
    c.drawImage(tmp.name, 0, 0, width=total_width * inch, height=total_height * inch)

    if spine_width > 0.5:
        c.setFillColorRGB(1, 1, 1)
        c.setFont("Helvetica-Bold", 10)
        c.saveState()
        c.translate(
            (bleed + trim_width + spine_width / 2) * inch,
            total_height / 2 * inch
        )
        c.rotate(90)
        c.drawCentredString(0, 0, f"{title}  |  {author}")
        c.restoreState()

    c.save()
    tmp.close()
    os.unlink(tmp.name)

    pdf_output.seek(0)
    return send_file(
        pdf_output,
        mimetype='application/pdf',
        as_attachment=True,
        download_name='book_cover.pdf'
    )

if __name__ == '__main__':
    app.run(debug=True)
